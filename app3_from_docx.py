# import the Document class 
# from the docx module 
from docx import Document 
import csv,re,os

# importing for XSD interface
from datetime import timedelta
from genericpath import isfile
import os
from flask import Flask, flash, redirect, request, send_file, url_for, session
from flask import render_template
from xml.dom import minidom
import pandas as pd
from werkzeug.utils import secure_filename
import csv
from pathlib import Path
from createXMLElements import *

# import for xsd interface ending here


folder = "docx_files/inhouse"
document_list = []
for filename in os.listdir(folder):
    if filename.endswith('.docx'):
        document_list.append(filename)
# create an instance of a 
# word document we want to open 
print(document_list)
for x in document_list:

	doc = Document(folder + "/" + x) 
	
	
	print('\nText in the 1st paragraph:->>>')  
	
	# for printing the complete document 
	print('\nThe whole content of the document:->>>\n') 
	no_of_tables = len(doc.tables)
	schema_filename = "schema_spec.csv"
	
	
	# below variables for writing to CSV files
	mylist = []
	my_list2 = []
	write_schema_data = False
	write_header_check_index = True
	check_first_record = True
	record_tables_index = 0
	loop_segment_start_check = 0
	# above variables for writing to CSV files
	
	
	#below code for xsd interface
	project_name = "Project_Name"
	
	schema_name = "Schema_Name"
	
	xmlobject = createXMLElements()
	xsd_root,header_sequence = xmlobject.createRoot(project_name,schema_name)
	xsd_loop_start_sequence = ''
	row = {}
	field_value = 0
	
	# Creating Root Segment Loop
	row['Type'] = 'LS' # Note : This is just added for Root Segment loop , code need to adjust for segment loops (inside root) also
	row['maxOccurs'] = '999999'
	row['minOccurs'] = '1'
	row['Loop/Rec Name'] = "Message_Root"
	atrributes = row
	xsd_loop_start,xsd_loop_start_sequence = xmlobject.createLoopSegment(header_sequence,xsd_loop_start_sequence,atrributes)
	#ending root segment Loop
	#above code for xsd interface
	
		
	
	
	first_record_segment_check_string = "Field Business Name"
	with open(folder + "/" + x + ".csv", 'w') as csvfile:
		for num in range(0, no_of_tables ):
			curr_table = doc.tables[num]
			
			for index,table_row in enumerate(curr_table.rows):
				# print(index)
				for index1, cell in enumerate(table_row.cells):
					if(index == 0 and first_record_segment_check_string in cell.text):
						write_schema_data = True
					else :
						if index == 0 and index1 == 0:
							write_schema_data = False
					if write_schema_data:
						if index == 0 and check_first_record :
							if index1 == 0 :
								mylist.append("Record_Name")            
							mylist.append(cell.text)
						else:
							if index > 0 :
								row['Type'] == 'FIELD'
								if index == 1 and index1 == 4:
									#below code is just to extract Record Name from Description column
									start_marker = "‘"
									end_marker = "’"
									if start_marker in cell.text:
										start_index = cell.text.find(start_marker) + len(start_marker)
										end_index = cell.text.find(end_marker, start_index)
										if end_index == -1:
											end_index = len(cell.text)
										extracted_string = cell.text[start_index:end_index]
										Record_Name = extracted_string
									else:
										Record_Name = str(re.sub("[^0-9]", "", cell.text))
									
									#above code is just to extract Record Name from Description column
								
								# data for Records and code for creating XSD Record Segment
								if index == 1 and index1 == 4:
									row['Type'] = "RC"
									row['Loop/Rec Name'] = Record_Name.replace(" ","_")
									row['maxOccurs'] = '10'
									row['minOccurs'] = '0'
									field_value = 1
									sequence = xmlobject.createRecordSegment(header_sequence,xsd_loop_start_sequence,atrributes)
								# data for Records and code for creating XSD Record Segment ends here
	
								if index1 == 0:
									mylist.append("ssss")
								mylist.append(cell.text)
				if write_schema_data:
					if index == 0 and check_first_record :
						# print(my_list2)
						writer = csv.writer(csvfile)
						writer.writerow(mylist)
						check_first_record = False
						del mylist[:]
					else:
						if index > 0 :
							#below for Fields inside Records
							pattern = re.compile(r'[^\w\s]|(?<!\S)_') # \W matches any non-alphanumeric character, _ matches underscore
	
							alphanumeric_string_Field_name = pattern.sub('', mylist[1].replace(" ","_"))
							row['Field Name'] = alphanumeric_string_Field_name
							row['Length'] = str(re.sub("[^0-9]", "", mylist[2]))
							row['maxOccurs'] = '1'
							if(mylist[3] == 'R'):
								row['minOccurs'] = '1'
							else:
								row['minOccurs'] = '0'
							if index == 1:
								default_value = Record_Name.replace(" ","_") 
								if default_value == "":
									default_value = mylist[1]
								default_value_length = row['Length']
								num_spaces = max(0, int(row['Length']) - len(default_value))
								formatted_string = '{:<{}}'.format(default_value, int(row['Length']))
								print(len(formatted_string))
								row['Default Value'] = formatted_string
							else:
								row['Default Value'] = ""
							xmlobject.createRecordFields(sequence,atrributes,index)
							field_value = field_value + 1
							# above code for Fields inside Records
							
							if Record_Name == "":
								mylist[0] = mylist[1]
							else:
								mylist[0] = Record_Name
							# print(my_list2)
							writer.writerow(mylist)
							del mylist[:]
	
	
	xml_str = (xsd_root.toprettyxml(indent="\t"))
	
	with open(folder + "/" + x.replace(" ", "_") + ".xsd", "w") as f:
			f.write(xml_str)
	
	